# src/aje_libs/bd/helpers/docx_helper.py
from docx import Document as DocxDocument
from typing import List, Dict, Optional
from ...common.logger import custom_logger

logger = custom_logger(__name__)

class DOCXHelper:
    """Helper para procesar archivos DOCX"""
    
    def __init__(self):
        pass
    
    def extract_text(self, file_path: str) -> str:
        """
        Extrae texto de un documento DOCX.
        
        :param file_path: Ruta al archivo DOCX
        :return: Texto extraído
        """
        try:
            doc = DocxDocument(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            extracted_text = "\n".join(paragraphs)
            logger.info(f"Texto extraído exitosamente de DOCX: {file_path}")
            return extracted_text
        except Exception as e:
            logger.error(f"Error extrayendo texto de DOCX: {e}")
            raise e
    
    def extract_paragraphs(self, file_path: str) -> List[Dict[str, str]]:
        """
        Extrae párrafos con su estilo.
        
        :param file_path: Ruta al archivo DOCX
        :return: Lista de párrafos con información de estilo
        """
        try:
            doc = DocxDocument(file_path)
            paragraphs = []
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    paragraph_info = {
                        'order': i + 1,
                        'text': paragraph.text,
                        'style': paragraph.style.name if paragraph.style else 'Normal'
                    }
                    paragraphs.append(paragraph_info)
            
            logger.info(f"Extraídos {len(paragraphs)} párrafos")
            return paragraphs
        except Exception as e:
            logger.error(f"Error extrayendo párrafos: {e}")
            raise e
    
    def extract_tables(self, file_path: str) -> List[List[List[str]]]:
        """
        Extrae tablas del documento.
        
        :param file_path: Ruta al archivo DOCX
        :return: Lista de tablas (cada tabla es una lista de filas)
        """
        try:
            doc = DocxDocument(file_path)
            tables_data = []
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            logger.info(f"Extraídas {len(tables_data)} tablas")
            return tables_data
        except Exception as e:
            logger.error(f"Error extrayendo tablas: {e}")
            raise e